import docx
d = docx.Document("c:\\Users\\User\\Desktop\\PROGRAMS\\python\\Pycharm Project Files\\documents\\word.docx")
print(d.paragraphs[0].text) # returns the first paragraph of the text
p = d.paragraphs[1]
print(p.runs[0].text)
p.runs[1].bold # returns True or False
p.runs[2].underline = True
p.runs[2].text = "changed text"
print(p.runs[2].text)
r = d.paragraphs[2]
r.style = "Heading 1"
d.add_paragraph("Hello, this is a pragraph l")
d.add_paragraph("New paragrph")
q = d.paragraphs[3]
q.add_run("Hello, this is a run")
q.runs[1].bold = True
d.save("c:\\Users\\User\\Desktop\\PROGRAMS\\python\\Pycharm Project Files\\documents\\word.docx")